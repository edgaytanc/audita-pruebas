import os
import openpyxl
from docx import Document

TAGS = [
    "[NOMBRE_AUDITOR]",
    "[FECHA_ELABORACION]",
    "[NOMBRE_SUPERVISOR]",
    "[FECHA_APROBACION]"
]

def check_excel(filepath):
    found_tags = set()
    try:
        wb = openpyxl.load_workbook(filepath, data_only=False)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for tag in TAGS:
                            if tag in cell.value:
                                found_tags.add(tag)
    except Exception as e:
        print(f"Error reading Excel {filepath}: {e}")
    return found_tags

def check_word(filepath):
    found_tags = set()
    try:
        doc = Document(filepath)
        # Check paragraphs
        for para in doc.paragraphs:
            for tag in TAGS:
                if tag in para.text:
                    found_tags.add(tag)
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for tag in TAGS:
                            if tag in para.text:
                                found_tags.add(tag)
        # Check headers/footers (simplified)
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        for tag in TAGS:
                            if tag in para.text:
                                found_tags.add(tag)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        for tag in TAGS:
                            if tag in para.text:
                                found_tags.add(tag)

    except Exception as e:
        print(f"Error reading Word {filepath}: {e}")
    return found_tags

def main():
    root_dir = "d:/temporal/audita-marcas"
    results = []

    print("Scanning files...")
    for root, dirs, files in os.walk(root_dir):
        for file in files:
            filepath = os.path.join(root, file)
            found = set()
            if file.endswith(".xlsx") and not file.startswith("~$"):
                found = check_excel(filepath)
            elif file.endswith(".docx") and not file.startswith("~$"):
                found = check_word(filepath)
            
            if found:
                results.append((filepath, list(found)))
                print(f"Found tags in {file}: {list(found)}")
            # else:
            #     print(f"No tags in {file}")

    print("\nSummary of files with tags:")
    for path, tags in results:
        print(f"{path}: {tags}")

if __name__ == "__main__":
    main()
