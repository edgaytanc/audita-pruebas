"""
Procesador de Marcas de Auditoría - Inyecta marcas en documentos descargados
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK
from openpyxl.styles import Font, PatternFill, Alignment
from django.db.models import Q
from auditoria.models import AuditMark
from auditoria.config.exclusion_rules import is_excluded
import re
import logging

logger = logging.getLogger(__name__)


class AuditMarkProcessor:
    """Procesa documentos para inyectar marcas de auditoría coincidentes."""

    def __init__(self, audit_id, filename):
        self.audit_id = audit_id
        self.filename = filename
        self.normalized_filename = self.normalize_text(filename)
        self.file_exclusion_info = is_excluded(filename)

    @staticmethod
    def normalize_text(text):
        """
        Normaliza texto para emparejamiento.

        Elimina números iniciales, extensiones de archivo, y caracteres especiales.
        Convierte a mayúsculas para comparación case-insensitive.
        """
        if not text:
            return ''

        # Eliminar números iniciales (incluyendo decimales) y espacios
        text = re.sub(r'^\d+(\.\d+)?\s*', '', text)

        # Eliminar extensiones de archivo
        text = re.sub(r'\.(xlsx|docx|xls|doc|pdf)$', '', text, flags=re.IGNORECASE)

        # Normalización: mayúsculas y solo alfanuméricos
        text = text.upper()
        text = re.sub(r'[^A-Z0-9]', '', text)

        return text

    def should_process_file(self):
        """Verifica si este archivo debe ser procesado para inyección de marcas."""
        if self.file_exclusion_info['is_excluded']:
            return (False, self.file_exclusion_info['reason'])
        return (True, None)

    def _is_mark_excluded(self, mark):
        """Verifica si una marca debe ser excluida del matching."""
        if not mark.work_paper_number:
            return (False, None)

        exclusion_info = is_excluded(mark.work_paper_number)
        if exclusion_info['is_excluded']:
            return (True, exclusion_info['reason'])

        return (False, None)

    def get_matching_marks(self):
        """Obtiene marcas que coincidan con el nombre de archivo."""
        marks = AuditMark.objects.filter(
            audit_id=self.audit_id,
            is_active=True
        ).exclude(
            Q(description__icontains='Ejemplo:') |
            Q(description__icontains='Example:')
        )

        matched_marks = []
        min_ratio = 0.8

        for mark in marks:
            if not mark.work_paper_number:
                continue

            is_excl, _ = self._is_mark_excluded(mark)
            if is_excl:
                continue

            normalized_wp = self.normalize_text(mark.work_paper_number)

            # Emparejamiento con ratio mínimo de similitud
            is_match = False

            if normalized_wp == self.normalized_filename:
                is_match = True
            elif normalized_wp in self.normalized_filename:
                ratio = len(normalized_wp) / len(self.normalized_filename)
                if ratio >= min_ratio:
                    is_match = True
            elif self.normalized_filename in normalized_wp:
                ratio = len(self.normalized_filename) / len(normalized_wp)
                if ratio >= min_ratio:
                    is_match = True

            if is_match:
                matched_marks.append(mark)

        return matched_marks

    def process_word_document(self, doc):
        """Agrega marcas de auditoría al documento Word."""
        try:
            should_process, _ = self.should_process_file()
            if not should_process:
                return doc

            matched_marks = self.get_matching_marks()
            if not matched_marks:
                return doc

            self._add_marks_to_word(doc, matched_marks)
            return doc

        except Exception as e:
            logger.error(f"Error procesando marcas para {self.filename}: {e}")
            return doc

    def _add_marks_to_word(self, doc, marks):
        """Agrega marcas en una nueva página al final del documento Word."""
        doc.add_page_break()
        doc.add_paragraph()

        # Título
        p_title = doc.add_paragraph()
        run_title = p_title.add_run('MARCAS DE AUDITORÍA UTILIZADAS:')
        run_title.bold = True
        run_title.font.size = Pt(14)
        run_title.font.color.rgb = RGBColor(192, 0, 0)

        doc.add_paragraph('─' * 80)

        # Marcas
        for mark in marks:
            p_mark = doc.add_paragraph()
            mark_text = f'{mark.symbol}    {mark.description}'
            run_mark = p_mark.add_run(mark_text)
            run_mark.font.size = Pt(11)
            run_mark.font.color.rgb = RGBColor(192, 0, 0)
            p_mark.paragraph_format.space_after = Pt(6)

    def process_excel_document(self, wb):
        """Agrega marcas de auditoría al final de la hoja Excel."""
        try:
            should_process, _ = self.should_process_file()
            if not should_process:
                return wb

            matched_marks = self.get_matching_marks()
            if not matched_marks:
                return wb

            self._add_marks_to_excel(wb, matched_marks)
            return wb

        except Exception as e:
            logger.error(f"Error procesando marcas para {self.filename}: {e}")
            return wb

    def _add_marks_to_excel(self, wb, marks):
        """Agrega sección de marcas al final de la hoja Excel."""
        ws = wb.active

        # Encontrar última fila con datos reales
        last_row = 0
        for row in range(ws.max_row, 0, -1):
            for col in range(1, ws.max_column + 1):
                cell_value = ws.cell(row=row, column=col).value
                if cell_value is not None and str(cell_value).strip():
                    last_row = row
                    break
            if last_row:
                break

        if not last_row:
            last_row = ws.max_row

        # Determinar columna de inicio basada en datos existentes
        start_column = 2
        col_counts = {}
        for row in range(1, min(21, ws.max_row + 1)):
            for col in range(1, min(10, ws.max_column + 1)):
                cell_value = ws.cell(row=row, column=col).value
                if cell_value is not None and str(cell_value).strip():
                    col_counts[col] = col_counts.get(col, 0) + 1
                    break

        if col_counts:
            start_column = max(col_counts, key=col_counts.get)
            if start_column < 2:
                start_column = 2

        start_row = last_row + 3

        # Título
        title_cell = ws.cell(row=start_row, column=start_column)
        title_cell.value = 'MARCAS DE AUDITORÍA UTILIZADAS:'
        title_cell.font = Font(bold=True, color='C00000', size=11)
        title_cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
        title_cell.alignment = Alignment(horizontal='left', vertical='center')

        # Marcas
        for i, mark in enumerate(marks, start=1):
            row_num = start_row + i
            cell = ws.cell(row=row_num, column=start_column)

            mark_text = f'{mark.symbol}  {mark.description}'
            if mark.symbol and mark.symbol[0] in ('=', '+', '-', '@'):
                mark_text = f"'{mark_text}"

            cell.value = mark_text
            cell.font = Font(bold=True, color='C00000', size=11)
            cell.fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
            cell.alignment = Alignment(horizontal='left', vertical='center')
