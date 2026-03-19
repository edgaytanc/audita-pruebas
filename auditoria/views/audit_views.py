"""
Vistas para gestión de auditorías.
Contiene funciones para listar auditorías financieras, internas y mostrar detalles.
"""

import os
import json
from .config import (
    render, redirect, HttpResponse, mark_safe, settings,
    Audit, login_required
)
from .utils import crear_mensaje_error, generar_html_estructura
from audits.decorators import audit_manager_required 
from django.contrib import messages
from django.http import JsonResponse, FileResponse
from django.template.loader import render_to_string
import mimetypes
try:
    from auditoria.utils.replacements_utils import get_replacements_config, build_replacements_dict
except ImportError:
    pass
try:
    from auditoria.utils.custom_replacements import build_custom_replacements
except ImportError:
    pass
from io import BytesIO
from docx import Document
from openpyxl import load_workbook
from auditoria.models import BalanceCuentas, RegistroAuxiliar, SaldoInicial
from datetime import datetime

@login_required
def auditorias_view(request):
    return render(request, 'auditoria/auditorias.html', {})

@login_required
def auditoria_financiera_view(request):
    user_role = request.user.role.name
    
    # Filtrado según el rol del usuario
    if user_role == "audit_manager":
        # Los administradores ven todas las auditorías que crearon
        auditorias_financieras = Audit.objects.filter(audit_manager=request.user, tipoAuditoria='F')
    else:
        # Los auditores regulares solo ven las auditorías asignadas a ellos
        auditorias_financieras = Audit.objects.filter(assigned_users=request.user, tipoAuditoria='F')
    
    if request.method == 'POST':
        audit_id = request.POST.get('audit_id')
        return redirect('auditoria_detalle', audit_id=audit_id)

    return render(request, 'auditoria/auditoria-financiera/auditoria_financiera.html', {
        'auditorias': auditorias_financieras
    })

@login_required
def auditoria_interna_view(request):
    user_role = request.user.role.name
    
    # Filtrado según el rol del usuario
    if user_role == "audit_manager":
        # Los administradores ven todas las auditorías que crearon
        auditorias_internas = Audit.objects.filter(audit_manager=request.user, tipoAuditoria='I')
    else:
        # Los auditores regulares solo ven las auditorías asignadas a ellos
        auditorias_internas = Audit.objects.filter(assigned_users=request.user, tipoAuditoria='I')
     
    if request.method == 'POST':
        audit_id = request.POST.get('audit_id')
        return redirect('auditoria_detalle', audit_id=audit_id)

    return render(request, 'auditoria/auditoria-interna/auditoria_interna.html', {
        'auditorias': auditorias_internas
    })


#Vista de revisoria    
@login_required
def auditoria_revisoria_view(request):
    user_role = request.user.role.name

    # Filtrado según el rol del usuario
    if user_role == "audit_manager":
        # Los administradores ven todas las auditorías de tipo Revisoría Fiscal que crearon
        auditorias_revisorias = Audit.objects.filter(audit_manager=request.user, tipoAuditoria='R')
    else:
        # Los auditores regulares solo ven las auditorías asignadas a ellos
        auditorias_revisorias = Audit.objects.filter(assigned_users=request.user, tipoAuditoria='R')

    # Redirección al detalle si el usuario selecciona una auditoría
    if request.method == 'POST':
        audit_id = request.POST.get('audit_id')
        return redirect('auditoria_detalle', audit_id=audit_id)

    # Renderizar la vista
    return render(request, 'auditoria/revisoria-fiscal/revisoria.html', {
        'auditorias': auditorias_revisorias
    })

@login_required
def auditoria_detalle_view(request, audit_id):
    user_role = request.user.role.name
    
    # Verificar primero si la auditoría existe
    try:
        if user_role == "audit_manager":
            # Los administradores solo pueden ver auditorías que ellos crearon
            audit = Audit.objects.get(id=audit_id, audit_manager=request.user)
        else:
            # Los auditores regulares solo pueden ver auditorías asignadas a ellos
            audit = Audit.objects.get(id=audit_id, assigned_users=request.user)
    except Audit.DoesNotExist:
        mensaje_error = crear_mensaje_error(
            "Auditoría no encontrada",
            "La auditoría solicitada no existe o no tienes permisos para acceder a ella."
        )
        return HttpResponse(mark_safe(mensaje_error), status=404)
    
    # Leer la estructura desde el archivo JSON correspondiente al tipo de auditoría
    # Leer la estructura desde el archivo JSON correspondiente al tipo de auditoría
    if audit.tipoAuditoria == 'I':
        json_path = os.path.join(settings.BASE_DIR, 'auditoria', 'config', 'folder_structure_interna.json')
    elif audit.tipoAuditoria == 'R':
        json_path = os.path.join(settings.BASE_DIR, 'auditoria', 'config', 'folder_structure_revisoria.json')
    else:
        json_path = os.path.join(settings.BASE_DIR, 'auditoria', 'config', 'folder_structure_financiera.json')
    
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            estructura_carpetas = json.load(f)
    except Exception:
        estructura_carpetas = {}  # Estructura vacía en caso de error
    
    # Usuario verificado si es administrador o tiene plan Mensual o Anual
    user_verified = request.user.username == 'administrador' or (hasattr(request.user, 'plan') and request.user.plan in ['M', 'A'])
    
    # Generar HTML de la estructura con la información de verificación
    estructura_html = generar_html_estructura(estructura_carpetas, audit_id, user_verified=user_verified)

    return render(request, 'auditoria/auditoria-detalle/auditoria_detalle.html', {
        'audit': audit,
        'estructura_html': mark_safe(estructura_html),  # Para que Django no escape el HTML
        'user_verified': user_verified,  # Pasar esta variable a la plantilla para debug
    })
# Función para crear el boton de descarga   
@login_required
@audit_manager_required
def generate_audit_document(req, audit_id):
    """
    Acción inicial del botón 'Generar documento'
    """
    messages.info(req, f"Generar documento de la auditoría {audit_id} (función en construcción).")
    return redirect("auditoria_detalle", audit_id=audit_id)



@login_required
def generar_documento_modal(request, audit_id):
    try:
        audit = Audit.objects.get(id=audit_id)
    except Audit.DoesNotExist:
        return JsonResponse({"error": "Auditoría no encontrada."}, status=404)

    # Auditor(es) desde assigned_users (exactamente como te funcionaba)
    auditores_qs = getattr(audit, "assigned_users", None)
    auditores = auditores_qs.all() if auditores_qs else []
    auditor_nombre = ", ".join([(u.get_full_name() or u.username) for u in auditores]) or "Sin asignar"

    # Supervisor desde audit_manager (exactamente como te funcionaba)
    supervisor_nombre = audit.audit_manager.get_full_name() if audit.audit_manager else "Sin asignar"

    html = render_to_string(
        "auditoria/modals/generar_documento_modal.html",
        {
            "audit": audit,
            "auditor_nombre": auditor_nombre,
            "supervisor_nombre": supervisor_nombre,
        },
        request=request,
    )
    return JsonResponse({"html": html})


@login_required
def obtener_carpetas_auditoria(request, audit_id):
    """
    Devuelve todas las carpetas y subcarpetas del tipo de auditoría (Financiera o Interna),
    leyendo desde static/templates_base_financiera o static/templates_base_interna.
    Se devuelven en orden descendente según su número inicial.
    """

    try:
        audit = Audit.objects.get(id=audit_id)
    except Audit.DoesNotExist:
        return JsonResponse({"error": "Auditoría no encontrada"}, status=404)

    # Selecciona la carpeta base según el tipo de auditoría
    if audit.tipoAuditoria == 'I':
        base_path = os.path.join(settings.BASE_DIR, 'static', 'templates_base_interna')
    else:
        base_path = os.path.join(settings.BASE_DIR, 'static', 'templates_base_financiera')

    if not os.path.exists(base_path):
        return JsonResponse({"error": f"No se encontró la carpeta base: {base_path}"}, status=404)

    carpetas = []

    # Recorrer carpetas y subcarpetas recursivamente
    for root, dirs, files in os.walk(base_path):
        # Obtener el nombre de la carpeta actual relativa a la base
        for d in dirs:
            relative_path = os.path.relpath(os.path.join(root, d), base_path)
            carpetas.append(relative_path.replace("\\", "/"))  # formato limpio compatible

    # Función auxiliar para extraer número de carpeta
    def extraer_numero(nombre):
        partes = nombre.split(" ", 1)
        try:
            return int(partes[0])
        except (ValueError, IndexError):
            return 0

    # Ordenar por número (descendente), luego alfabéticamente
    carpetas.sort(key=lambda x: (extraer_numero(x), x.lower()), reverse=False)

    return JsonResponse({"carpetas": carpetas})

@login_required
def obtener_documentos_auditoria(request, audit_id, carpeta_nombre):
    """
    Devuelve los documentos (.docx y .xlsx) dentro de una carpeta seleccionada,
    ordenados numéricamente por nombre.
    """
    from urllib.parse import unquote
    import re

    try:
        audit = Audit.objects.get(id=audit_id)
    except Audit.DoesNotExist:
        return JsonResponse({"error": "Auditoría no encontrada"}, status=404)

    carpeta_nombre = unquote(carpeta_nombre)

    if audit.tipoAuditoria == 'I':
        base_path = os.path.join(settings.BASE_DIR, 'static', 'templates_base_interna')
    else:
        base_path = os.path.join(settings.BASE_DIR, 'static', 'templates_base_financiera')

    carpeta_path = os.path.join(base_path, carpeta_nombre)
    if not os.path.exists(carpeta_path):
        return JsonResponse({"error": f"No se encontró la carpeta: {carpeta_nombre}"}, status=404)

    archivos = [
        f for f in os.listdir(carpeta_path)
        if os.path.isfile(os.path.join(carpeta_path, f)) and f.lower().endswith(('.docx', '.xlsx'))
    ]

    # 🔢 Ordenar numéricamente por prefijo
    def num_key(nombre):
        match = re.match(r'^(\d+(?:\.\d+)?)', nombre)
        return float(match.group(1)) if match else float('inf')

    archivos.sort(key=num_key)

    return JsonResponse({"documentos": archivos})


# ==========================================================
# 🔸 Sub-funciones auxiliares
# ==========================================================

def _get_audit_object(audit_id):
    """Obtiene la auditoría o devuelve error JSON."""
    try:
        return Audit.objects.get(id=audit_id)
    except Audit.DoesNotExist:
        return None


def _resolve_template_path(audit, carpeta, documento):
    """Devuelve la ruta absoluta del archivo base (Word o Excel)."""
    base_dir = (
        "templates_base_interna"
        if audit.tipoAuditoria == "I"
        else "templates_base_financiera"
    )
    return os.path.join(settings.BASE_DIR, "static", base_dir, carpeta, documento)


def _apply_word_replacements(file_path, replacements):
    """Abre el archivo Word, aplica los reemplazos y devuelve un buffer."""
    doc = Document(file_path)

    # Reemplazo en párrafos
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# ==========================================================
# 🔸 FUNCIÓN PRINCIPAL
# ==========================================================
@login_required
def descargar_documento_auditoria(request, audit_id):
    """Descarga el documento ya automatizado con reemplazos, datos importados y etiquetas personalizadas."""

    from auditoria.word_utils import modify_document_word  # ✅ usamos su automatización oficial

    # -------------------- helpers SOLO para fechas --------------------
    # (no tocan lo que ya funciona; actúan después como “pasada extra”)
    from docx.oxml.ns import qn

    def _replace_wt_spanning_tokens(container, token: str, value: str):
        """
        Reemplaza `token` aunque esté dividido en varios <w:t>.
        Conserva el formato fuera del token. El texto resultante del token
        quedará con el formato del primer <w:t> del grupo (lo típico en firmas).
        """
        t_nodes = container._element.xpath('.//w:t')
        if not t_nodes:
            return

        # 1) Reemplazos directos dentro de un solo w:t
        for t in t_nodes:
            if t.text and token in t.text:
                t.text = t.text.replace(token, value)

        # 2) Si aún queda el token, probar en ventanas de 2 y 3 nodos contiguos
        #    (suele ser el caso más común en estas plantillas)
        def replace_in_window(win_size: int):
            i = 0
            while i <= len(t_nodes) - win_size:
                parts = [(t_nodes[i+k].text or '') for k in range(win_size)]
                joined = ''.join(parts)
                if token in joined:
                    new_joined = joined.replace(token, value)

                    # Redistribuir el nuevo texto en los mismos nodos (manteniendo estructura)
                    remaining = new_joined
                    for k in range(win_size):
                        original = parts[k]
                        take = min(len(original), len(remaining))
                        t_nodes[i+k].text = remaining[:take]
                        remaining = remaining[take:]
                    if remaining:
                        # si sobró (pasa cuando value es más largo que el token original),
                        # lo añadimos al último nodo de la ventana
                        t_nodes[i+win_size-1].text = (t_nodes[i+win_size-1].text or '') + remaining
                i += 1

        # ventanas de 2 y 3 nodos
        replace_in_window(2)
        replace_in_window(3)



    # -----------------------------------------------------------------

    # Obtener la auditoría
    audit = _get_audit_object(audit_id)
    if not audit:
        return JsonResponse({"error": "Auditoría no encontrada."}, status=404)

    carpeta = request.GET.get("carpeta")
    documento = request.GET.get("documento")

    if not carpeta or not documento:
        return JsonResponse({"error": "Faltan parámetros."}, status=400)

    # Determinar ruta
    file_path = _resolve_template_path(audit, carpeta, documento)
    if not os.path.exists(file_path):
        return JsonResponse({"error": "El archivo no existe."}, status=404)

    # --- Fechas ---
    def _fmt_ddmmyyyy(date_str):
        if not date_str:
            return ""
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d-%m-%Y")
        except Exception:
            return date_str

    fecha_elaboracion = _fmt_ddmmyyyy(request.GET.get("fecha_elaboracion"))
    fecha_aprobacion  = _fmt_ddmmyyyy(request.GET.get("fecha_aprobacion"))

    # --- Reemplazos base ---
    replacements_config = get_replacements_config()
    fecha_inicio = audit.fechaInit.strftime("%d-%m-%Y") if audit.fechaInit else ""
    fecha_fin    = audit.fechaEnd.strftime("%d-%m-%Y") if audit.fechaEnd else ""
    replacements = build_replacements_dict(replacements_config, audit, fecha_inicio, fecha_fin)

    # --- Reemplazos personalizados tuyos ---
    auditores_qs = getattr(audit, "assigned_users", None)
    auditores = auditores_qs.all() if auditores_qs else []
    auditor_nombre = ", ".join([(u.get_full_name() or u.username) for u in auditores]) or "Sin asignar"
    supervisor_nombre = audit.audit_manager.get_full_name() if audit.audit_manager else "Sin asignar"

    replacements["[NOMBRE_AUDITOR]"]    = auditor_nombre
    replacements["[NOMBRE_SUPERVISOR]"] = supervisor_nombre
    replacements["[FECHA_ELABORACION]"] = fecha_elaboracion or ""
    replacements["[FECHA_APROBACION]"]  = fecha_aprobacion or ""

    # --- Detectar tipo de archivo ---
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        try:
            # 🧩 Paso 1: automatización oficial (tu lógica actual)
            doc = modify_document_word(file_path, audit)

            # 🧩 Paso 2: aplicar tus reemplazos (igual que tenías, preservando formato)
            # EXCEPTO los que ahora haremos robustos: fechas Y nombres
            ROBUST_TAGS = ["[FECHA_ELABORACION]", "[FECHA_APROBACION]", "[NOMBRE_AUDITOR]", "[NOMBRE_SUPERVISOR]"]
            
            replacements_filtrados = {
                k: v for k, v in replacements.items()
                if (not any(header in k for header in ["Entidad:", "Auditoría:", "Período:"]))
                and (k not in ROBUST_TAGS)
            }

            # Reemplazo por run (lo que ya tienes y funciona para el resto)
            for p in doc.paragraphs:
                for run in p.runs:
                    for key, value in replacements_filtrados.items():
                        if key in run.text:
                            run.text = run.text.replace(key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                for key, value in replacements_filtrados.items():
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

            # 🧩 Paso 3: **pasada robusta PARA TODOS LOS TAGS CRÍTICOS** (cubre tokens partidos en varios w:t)
            def _replace_tags_everywhere(doc, tags_map: dict):
                """
                Aplica _replace_wt_spanning_tokens a todo el documento para las etiquetas dadas.
                """
                # Párrafos
                for p in doc.paragraphs:
                    for tag, val in tags_map.items():
                        # Chequeo rápido antes de procesar XML
                        if tag in ''.join(run.text or '' for run in p.runs) or any(tag in (t.text or '') for t in p._element.xpath('.//w:t')):
                            _replace_wt_spanning_tokens(p, tag, val)

                # Tablas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for tag, val in tags_map.items():
                                    if tag in ''.join(run.text or '' for run in p.runs) or any(tag in (t.text or '') for t in p._element.xpath('.//w:t')):
                                        _replace_wt_spanning_tokens(p, tag, val)

            _replace_tags_everywhere(doc, {
                "[FECHA_ELABORACION]": replacements.get("[FECHA_ELABORACION]", ""),
                "[FECHA_APROBACION]":  replacements.get("[FECHA_APROBACION]",  ""),
                "[NOMBRE_AUDITOR]":    replacements.get("[NOMBRE_AUDITOR]",    ""),
                "[NOMBRE_SUPERVISOR]": replacements.get("[NOMBRE_SUPERVISOR]", ""),
            })

            # 🧩 Paso 4: guardar en buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        except Exception as e:
            import traceback
            traceback.print_exc()
            return JsonResponse({"error": f"Error procesando el documento Word: {str(e)}"}, status=500)

    elif ext in [".xlsx", ".xlsm"]:
        # ✅ Excel: lo de siempre (sin cambios)
        from auditoria.views.download_views import modify_document_excel, modify_document_excel_with_macros

        if ext == ".xlsm":
            processed_path = modify_document_excel_with_macros(file_path, audit)
            wb = load_workbook(processed_path, keep_vba=True)

            for sh in wb.worksheets:
                for row in sh.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            v = cell.value
                            for k, val in replacements.items():
                                if k in v:
                                    v = v.replace(k, val)
                            cell.value = v

            bio = BytesIO()
            wb.save(bio)
            bio.seek(0)
            buffer = bio
            content_type = "application/vnd.ms-excel.sheet.macroEnabled.12"

        else:
            wb = modify_document_excel(file_path, audit)
            for sh in wb.worksheets:
                for row in sh.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str):
                            v = cell.value
                            for k, val in replacements.items():
                                if k in v:
                                    v = v.replace(k, val)
                            cell.value = v

            bio = BytesIO()
            wb.save(bio)
            bio.seek(0)
            buffer = bio
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    else:
        return JsonResponse({"error": "Formato de archivo no soportado."}, status=400)

    # --- Respuesta final ---
    response = HttpResponse(buffer, content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{documento}"'
    return response
