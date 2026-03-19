"""
Test para verificar que las marcas de auditoría aparecen SOLO en la última página.

Bug fix: Anteriormente las marcas aparecían en todas las páginas.
Ahora deben aparecer solo en la última página.
"""

from django.test import TestCase
from django.contrib.auth import get_user_model
from django.utils import timezone
from audits.models import Audit
from users.models import Roles
from auditoria.models import AuditMark
from auditoria.services.audit_mark_processor import AuditMarkProcessor
from docx import Document
from docx.oxml.shared import OxmlElement, qn

User = get_user_model()


class WordFooterLastPageTestCase(TestCase):
    """
    Test para verificar que las marcas se agregan SOLO a la última página
    """

    def setUp(self):
        """Configuración inicial"""
        self.role = Roles.objects.create(name='audit_manager')
        self.user = User.objects.create_user(
            username='test_footer',
            email='footer@test.com',
            password='testpass123',
            role=self.role
        )

        self.audit = Audit.objects.create(
            title="Test Footer Fix",
            identidad="TEST-FOOTER-2025",
            fechaInit=timezone.now(),
            fechaEnd=timezone.now(),
            tipoAuditoria="F",
            audit_manager=self.user
        )

    def test_marks_appear_only_on_last_page(self):
        """
        Test: Marcas deben aparecer SOLO en la última página, no en todas
        NUEVA IMPLEMENTACIÓN: Las marcas se agregan en el cuerpo de una nueva página final
        """
        # Crear marcas de prueba
        mark1 = AuditMark.objects.create(
            audit=self.audit,
            symbol='✓',
            description='Verificado contra fuente original',
            work_paper_number='TEST DOCUMENTO',
            is_active=True
        )

        mark2 = AuditMark.objects.create(
            audit=self.audit,
            symbol='∑',
            description='Suma verificada',
            work_paper_number='TEST DOCUMENTO',
            is_active=True
        )

        # Crear documento Word con contenido inicial
        doc = Document()

        # Agregar contenido inicial (simulando documento existente)
        doc.add_paragraph("Contenido de la página 1")
        doc.add_paragraph("Más contenido...")

        # Contar párrafos antes de procesar
        paragraphs_before = len(doc.paragraphs)

        # Procesar documento con AuditMarkProcessor
        processor = AuditMarkProcessor(self.audit.id, 'TEST DOCUMENTO.docx')

        # Llamar directamente al método que agrega marcas
        matched_marks = [mark1, mark2]
        processor._add_marks_to_footer(doc, matched_marks)

        # ═══════════════════════════════════════════════════════════
        # VERIFICACIÓN CRÍTICA: Marcas agregadas al final del documento
        # ═══════════════════════════════════════════════════════════

        # Debe haber más párrafos ahora (salto de página + título + línea + marcas)
        paragraphs_after = len(doc.paragraphs)
        self.assertGreater(paragraphs_after, paragraphs_before,
                          "ERROR: No se agregaron párrafos para las marcas")

        # Obtener todo el texto del documento
        full_text = '\n'.join([p.text for p in doc.paragraphs])

        # Verificar que las marcas están en el documento
        self.assertIn('MARCAS DE AUDITORÍA UTILIZADAS', full_text,
                     "ERROR: El título de marcas no aparece en el documento")
        self.assertIn('✓', full_text,
                     "ERROR: El símbolo ✓ no está en el documento")
        self.assertIn('Verificado contra fuente original', full_text,
                     "ERROR: La descripción no está en el documento")
        self.assertIn('∑', full_text,
                     "ERROR: El símbolo ∑ no está en el documento")
        self.assertIn('Suma verificada', full_text,
                     "ERROR: La descripción no está en el documento")

        # Verificar que las marcas están al FINAL (en los últimos párrafos)
        last_10_paragraphs = '\n'.join([p.text for p in doc.paragraphs[-10:]])
        self.assertIn('MARCAS DE AUDITORÍA UTILIZADAS', last_10_paragraphs,
                     "ERROR: Las marcas no están al final del documento")

        # Verificar que las marcas NO están en los primeros párrafos (contenido original)
        first_2_paragraphs = '\n'.join([p.text for p in doc.paragraphs[:2]])
        self.assertNotIn('MARCAS DE AUDITORÍA UTILIZADAS', first_2_paragraphs,
                        "ERROR: Las marcas aparecen al principio del documento")

    def test_single_section_document(self):
        """
        Test: Documento con una sola sección (una página) debe tener marcas al final
        NUEVA IMPLEMENTACIÓN: Las marcas se agregan en el cuerpo de una nueva página
        """
        # Crear marca de prueba
        mark = AuditMark.objects.create(
            audit=self.audit,
            symbol='©',
            description='Comprobado',
            work_paper_number='SINGLE PAGE',
            is_active=True
        )

        # Crear documento con una sola página
        doc = Document()
        doc.add_paragraph("Contenido de la única página")

        # Procesar documento
        processor = AuditMarkProcessor(self.audit.id, 'SINGLE PAGE.docx')
        processor._add_marks_to_footer(doc, [mark])

        # Verificar que las marcas están en el documento
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        self.assertIn('MARCAS DE AUDITORÍA UTILIZADAS', full_text)
        self.assertIn('©', full_text)
        self.assertIn('Comprobado', full_text)

    def test_empty_document(self):
        """
        Test: Documento sin secciones no debe causar error
        """
        # Crear documento vacío (edge case)
        doc = Document()

        # Verificar que tiene al menos 1 sección por defecto
        self.assertGreater(len(doc.sections), 0)

        # Procesar no debe causar error
        processor = AuditMarkProcessor(self.audit.id, 'EMPTY.docx')

        try:
            processor._add_marks_to_footer(doc, [])
            # Si llegamos aquí, no hubo error
            self.assertTrue(True)
        except Exception as e:
            self.fail(f"El procesamiento de documento vacío causó error: {e}")
